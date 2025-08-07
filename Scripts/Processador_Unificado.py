# Scripts/Processador_Unificado.py
# VERSÃO FINAL 2.1 - Inclui a lógica aprimorada de etiquetagem de origem.

import os
import sys
import time
import json
import re
import difflib
import openai
import tiktoken
import requests
import base64
import io
import traceback
from docx import Document
from openpyxl import Workbook
from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from dotenv import load_dotenv
from typing import Optional, List, Dict, Any

# --- Bloco de Inicialização Django ---
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.append(str(PROJECT_ROOT))
load_dotenv(PROJECT_ROOT / ".env")
import django
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "dossel_project.settings")
django.setup()
import argparse
from revisor.models import QueueEntry
from django.conf import settings

# --- Argumentos e Configurações Globais ---
parser = argparse.ArgumentParser(description="Processador Unificado de Documentos Dossel.")
parser.add_argument("arquivo", help="Nome do arquivo DOCX (com extensão) que está na PASTA_ENTRADA.")
parser.add_argument("--id_revisao", type=int, required=True, help="ID da revisão, passado pelo Gerenciador.")
args = parser.parse_args()

client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
ENCODER = tiktoken.encoding_for_model("gpt-4o")
MAX_WORKERS = 5
TIMEOUT_API = 300
MAX_RETRIES = 2

PASTA_ENTRADA = Path(settings.PASTA_ENTRADA)
PASTA_SAIDA = Path(settings.PASTA_SAIDA)

VALOR_INPUT = 0.005
VALOR_OUTPUT = 0.015
COTACAO_DOLAR = 5.65

# Desativa o TQDM se a saída não for um terminal interativo (ex: log de arquivo)
DISABLE_TQDM = not sys.stdout.isatty()

# --- Prompts para a IA ---
PROMPT_CATEGORIZACAO = "Você é um classificador de texto. Analise o trecho e retorne APENAS UMA das seguintes categorias com base no erro mais evidente: 'textual' (erros de gramática, ortografia, concordância, pontuação), 'bibliografico' (referências, citações ABNT), 'logico' (clareza, coesão, fluidez), ou 'nenhum' se não houver erros. Se for um título, sumário ou legenda, retorne 'nenhum'."
PROMPT_REVISAO_TEXTUAL = "Você é um revisor de textos acadêmicos e técnicos, especialista na norma culta da língua portuguesa. Sua tarefa é corrigir o trecho abaixo com o máximo rigor. Foco em erros de gramática, ortografia, concordância verbal e nominal, regência, crase e pontuação. NÃO altere o estilo de escrita, a voz do autor ou a terminologia técnica. Se o texto já estiver correto, repita-o integralmente. Responda APENAS com o texto corrigido."
PROMPT_REVISAO_BIBLIO = "Você é um especialista em ABNT NBR 6023. Revise e padronize a referência bibliográfica abaixo. Responda APENAS com a referência corrigida."
PROMPT_JUSTIFICATIVA = "Você é um revisor. Compare o 'Texto Original' com o 'Texto Corrigido' e explique de forma concisa e técnica, em uma única frase, qual foi a principal correção aplicada.\n\nTexto Original: \"{texto_original}\"\nTexto Corrigido: \"{texto_corrigido}\""
PROMPT_EXTRAIR_FATOS = "Sua tarefa é ler o parágrafo a seguir e extrair apenas as alegações factuais, dados numéricos, ou definições chave. Seja extremamente conciso. Se o parágrafo for subjetivo ou não contiver fatos, responda com 'N/A'.\n\nParágrafo: \"{texto}\""
# --- NOVOS PROMPTS DE ANÁLISE AVANÇADA ---
PROMPT_ANALISE_CRITICA_LOCAL = "Sua tarefa é atuar como um analista crítico e lógico. Leia o trecho a seguir e: 1. Extraia concisamente as alegações factuais e dados chave. 2. Aponte explicitamente QUALQUER inconsistência lógica, contradição, dado impossível (ex: '101%'), ou saltos lógicos onde uma conclusão não deriva claramente das premissas. Se não houver nada a relatar, responda 'N/A'.\n\nTrecho: \"{texto}\""
PROMPT_VALIDACAO_TABELA_ARITMETICA = "Você é um analista de dados. A seguir está o conteúdo de uma tabela. Sua única tarefa é verificar se os cálculos matemáticos na tabela estão corretos, especialmente em colunas de 'Diferença' ou 'Total'. Se encontrar algum erro aritmético, liste-o de forma clara, especificando a linha e o cálculo incorreto. Se todos os cálculos estiverem corretos, responda com 'Cálculos da tabela validados.'\n\nConteúdo da Tabela:\n{tabela_em_texto}"
PROMPT_DESCREVER_IMAGEM = "Sua tarefa é fornecer uma descrição objetiva e concisa dos principais elementos na imagem. Use o texto da legenda fornecida apenas como contexto. NÃO inclua nenhum prefixo como 'Figura X:'. Responda APENAS com a descrição do conteúdo visual.\n\nLegenda para Contexto: '{legenda}'"
PROMPT_ANALISE_GLOBAL_DOSSIE= "Você é um analista lógico. O dossiê abaixo contém uma lista de fatos, dados e descrições extraídos de um documento, incluindo validações automáticas. Sua tarefa é encontrar e listar TODAS as contradições lógicas ou factuais entre os itens desta lista. Se, e somente se, a lista de contradições estiver vazia, responda com a frase 'Nenhuma inconsistência global foi encontrada.'.\n\nDossiê para Análise:\n\n{dossie}"



# --- Funções Auxiliares e de Validação ---
def contar_tokens(txt: str) -> int: return len(ENCODER.encode(txt))
def similaridade(a: str, b: str) -> float: return difflib.SequenceMatcher(None, a.lower(), b.lower()).ratio()

def _br2float(valor: str) -> Optional[float]:
    if valor is None: return None
    s = str(valor).strip().replace(".", "").replace(" ", "").replace(",", ".")
    s = re.sub(r"[^0-9\.\-]", "", s)
    if s in ("", "-", ".", "-."): return None
    try: return float(s)
    except ValueError: return None
def validar_paragrafo_textual(texto: str) -> List[Dict[str, Any]]:
    issues = []
    # ... (implementação completa das suas funções de validação aqui) ...
    return issues
def validar_tabela_diferenca(linhas_tabela: List[List[str]]) -> List[Dict[str, Any]]:
    inconsist = []
    # ... (implementação completa das suas funções de validação aqui) ...
    return inconsist

def executar_chamada_api(prompt_sistema: str, trecho_usuario: str):
    tokens_in, tokens_out = 0, 0
    for _ in range(MAX_RETRIES):
        try:
            response = client.chat.completions.create(
                model="gpt-4o", messages=[{"role": "system", "content": prompt_sistema}, {"role": "user", "content": trecho_usuario}],
                temperature=0.1, timeout=TIMEOUT_API
            )
            resposta_texto = response.choices[0].message.content.strip()
            tokens_in = response.usage.prompt_tokens
            tokens_out = response.usage.completion_tokens
            return resposta_texto, tokens_in, tokens_out
        except Exception as e:
            print(f"  Aviso: Erro na API OpenAI ({e}). Tentando novamente...")
            time.sleep(2)
    return None, tokens_in, tokens_out

def validar_doi(doi):
    if not doi: return False
    try:
        r = requests.head(f"https://doi.org/{doi}", timeout=7, allow_redirects=True)
        return r.status_code == 200
    except requests.RequestException: return False

def validar_isbn(texto):
    return bool(re.search(r"\b(?:ISBN(?:-1[03])?:? )?(?=[0-9X]{10}$|(?=(?:[0-9]+[- ]){3})[- 0-9X]{13}$|97[89][0-9]{10}$|(?=(?:[0-9]+[- ]){4})[- 0-9]{17}$)(?:97[89][- ]?)?[0-9]{1,5}[- ]?[0-9]+[- ]?[0-9]+[- ]?[0-9X]\b", texto))

def validar_url(texto):
    urls = re.findall(r'https?://[^\s<>"\'`]+', texto)
    if not urls: return False
    for u in urls:
        try:
            r = requests.head(u, timeout=7, allow_redirects=True)
            if r.status_code < 400: return True
        except requests.RequestException: continue
    return False

def extrair_imagens_do_bloco(paragrafo, doc):
    imagens = []
    for run in paragrafo.runs:
        for inline_shape in run.element.xpath('.//wp:inline'):
            rId = inline_shape.xpath('.//a:blip/@r:embed')[0]
            image_part = doc.part.related_parts[rId]
            imagens.append(image_part.blob)
    return imagens

def analisar_imagem_com_visao(blob_imagem, legenda):
    try:
        prompt_texto = PROMPT_DESCREVER_IMAGEM.format(legenda=legenda or "Nenhuma")
        response = client.chat.completions.create(
            model="gpt-4o", messages=[{"role": "user", "content": [{"type": "text", "text": prompt_texto}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64.b64encode(blob_imagem).decode('utf-8')}"}}]}], max_tokens=300,
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"  Aviso: Falha ao analisar imagem. Erro: {e}")
        return "Falha ao analisar a imagem."

# --- Funções do Pipeline ---
# Scripts/Processador_Unificado.py -> DENTRO DO SEU SCRIPT

def analisar_tabela_para_erros_aritmeticos(table, executar_chamada_api_func):
    """Converte uma tabela docx em texto e usa a IA para verificar erros de cálculo."""
    try:
        tabela_em_texto = ""
        for row in table.rows:
            celulas = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            tabela_em_texto += "| " + " | ".join(celulas) + " |\n"
        
        if len(tabela_em_texto) < 50: return None

        prompt = PROMPT_VALIDACAO_TABELA_ARITMETICA.format(tabela_em_texto=tabela_em_texto)
        resultado, ti, to = executar_chamada_api_func("", prompt)

        if resultado and "validados" not in resultado.lower():
            return f"Possível erro aritmético na tabela: {resultado}", ti, to
    except Exception as e:
        print(f"  Aviso: Falha ao analisar tabela para erros aritméticos. Erro: {e}")
    return None, 0, 0

def mapear_conteudo_do_documento(doc: Document):
    print("  FASE 0: Mapeando conteúdo e etiquetando origens (versão completa)...")
    itens_mapeados = []
    paragrafos_corpo = doc.paragraphs
    regex_legenda = re.compile(r"^(figura|fig|foto|imagem)[\s-]?\d+.*", re.IGNORECASE)
    
    # --- Parte 1: Percorre o corpo principal para textos e pares imagem-legenda ---
    p_idx = 0
    i = 0
    while i < len(paragrafos_corpo):
        p_idx += 1
        p = paragrafos_corpo[i]
        texto_paragrafo = p.text.strip()
        
        imagens_blob = extrair_imagens_do_bloco(p, doc)

        if imagens_blob:
            legenda_texto, legenda_rotulo = "", ""
            if i + 1 < len(paragrafos_corpo) and regex_legenda.match(paragrafos_corpo[i+1].text.strip()):
                legenda_texto = paragrafos_corpo[i+1].text.strip()
                legenda_rotulo = legenda_texto
                i += 1
            
            for blob in imagens_blob:
                origem = legenda_rotulo or f"Imagem próxima ao Parágrafo {p_idx}"
                itens_mapeados.append({
                    "tipo": "imagem", "blob": blob,
                    "legenda_texto": legenda_texto, "origem": origem, "objeto": p
                })
        
        elif texto_paragrafo:
            itens_mapeados.append({"tipo": "texto", "texto": texto_paragrafo, "origem": f"Parágrafo {p_idx}", "objeto": p})
        
        i += 1
        
    # --- Parte 2: Percorre todas as tabelas ---
    print("  Mapeando conteúdo de tabelas...")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            # A LINHA ABAIXO FOI CORRIGIDA
            for c_idx, cell in enumerate(row.cells): # CORREÇÃO: iterar sobre row.cells, não table.cells
                for p_cell in cell.paragraphs:
                    if p_cell.text.strip():
                        origem = f"Tabela {t_idx+1}, Linha {r_idx+1}, Coluna {c_idx+1}"
                        itens_mapeados.append({"tipo": "texto", "texto": p_cell.text.strip(), "origem": origem, "objeto": p_cell})

    # --- Parte 3: Percorre cabeçalhos e rodapés ---
    print("  Mapeando conteúdo de cabeçalhos e rodapés...")
    for s_idx, section in enumerate(doc.sections):
        for p_header in section.header.paragraphs:
            if p_header.text.strip():
                origem = f"Cabeçalho da Seção {s_idx+1}"
                itens_mapeados.append({"tipo": "texto", "texto": p_header.text.strip(), "origem": origem, "objeto": p_header})
        for p_footer in section.footer.paragraphs:
            if p_footer.text.strip():
                origem = f"Rodapé da Seção {s_idx+1}"
                itens_mapeados.append({"tipo": "texto", "texto": p_footer.text.strip(), "origem": origem, "objeto": p_footer})

    print(f"  Mapeamento concluído. {len(itens_mapeados)} itens de conteúdo etiquetados (incluindo tabelas, cabeçalhos e rodapés).")
    return itens_mapeados

def fase_1_mapeamento_e_categorizacao(itens: list):
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        future_map = {executor.submit(executar_chamada_api, "", PROMPT_CATEGORIZACAO.format(texto=item["texto"])): item for item in itens if item["tipo"] == "texto"}
        progresso = tqdm(as_completed(future_map), total=len(future_map), desc="FASE 1: Categorizando", disable=DISABLE_TQDM)
        for future in progresso:
            item_original = future_map[future]
            categoria, _, _ = future.result()
            item_original["categoria"] = categoria.lower() if categoria else "nenhum"
    return itens

def fase_2_revisao_e_correcao(itens: list):
    metricas = {"textual": {"in": 0, "out": 0}, "bibliografico": {"in": 0, "out": 0}, "justificativa": {"in": 0, "out": 0}}
    revisoes_log = []
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        future_map = {}
        for item in itens:
            if item["tipo"] == "texto":
                prompt_revisao, tipo_revisao = (None, None)
                if item.get("categoria") == "textual": prompt_revisao, tipo_revisao = (PROMPT_REVISAO_TEXTUAL, "textual")
                elif item.get("categoria") == "bibliografico": prompt_revisao, tipo_revisao = (PROMPT_REVISAO_BIBLIO, "bibliografico")
                if prompt_revisao:
                    future = executor.submit(executar_chamada_api, prompt_revisao, item["texto"])
                    future_map[future] = (item, tipo_revisao)
        progresso = tqdm(as_completed(future_map), total=len(future_map), desc="FASE 2: Revisando", disable=DISABLE_TQDM)
        for future in progresso:
            item_original, tipo_revisao = future_map[future]
            texto_corrigido, ti, to = future.result()
            metricas[tipo_revisao]["in"] += ti
            metricas[tipo_revisao]["out"] += to
            if texto_corrigido and similaridade(item_original["texto"], texto_corrigido) < 0.999:
                item_original["objeto"].text = texto_corrigido
                item_original["texto_corrigido"] = texto_corrigido
                justificativa, ti_just, to_just = executar_chamada_api("", PROMPT_JUSTIFICATIVA.format(texto_original=item_original["texto"], texto_corrigido=texto_corrigido))
                metricas["justificativa"]["in"] += ti_just
                metricas["justificativa"]["out"] += to_just
                comentario_validacao = ""
                if tipo_revisao == "bibliografico":
                    doi_match = re.search(r"10\.\d{4,9}/[-._;()/:A-Z0-9]+", texto_corrigido, re.I)
                    obs = [f"DOI {'Válido' if validar_doi(doi_match.group(0) if doi_match else None) else 'Inválido/Não Encontrado'}", f"ISBN {'Encontrado' if validar_isbn(texto_corrigido) else 'Não Encontrado'}", f"URL {'Acessível' if validar_url(texto_corrigido) else 'Quebrada/Não Encontrada'}"]
                    comentario_validacao = " | ".join(obs)
                revisoes_log.append({"tipo": tipo_revisao, "original": item_original["texto"], "corrigido": texto_corrigido, "validacao": comentario_validacao, "justificativa": justificativa or "Correção aplicada."})
    return itens, metricas, revisoes_log

def fase_3_verificacao_cruzada(itens: list):
    print("  FASE 3: Verificação Cruzada de Citações iniciada.")
    regex_citacao = r'\(([A-Z][A-Z\s,;]+?,\s*\d{4}[a-z]?)\)'
    citacoes_no_texto, referencias_bibliograficas, inconsistencias = set(), [], []
    for item in itens:
        if item["tipo"] == "texto":
            texto_analise = item.get("texto_corrigido", item["texto"])
            if item.get("categoria") == "bibliografico": referencias_bibliograficas.append(texto_analise.lower())
            else:
                for match in re.findall(regex_citacao, texto_analise.upper()): citacoes_no_texto.add(match)
    for citacao in sorted(list(citacoes_no_texto)):
        autor, ano = citacao.split(',')[0].strip(), citacao.split(',')[1].strip()
        if not any(autor.lower() in ref and ano in ref for ref in referencias_bibliograficas):
            inconsistencias.append(f"A citação ({citacao}) parece não ter uma referência correspondente na bibliografia.")
    return inconsistencias

def fase_4_analise_multimodal_e_global(itens: list, doc: Document):
    print("\n--- INICIANDO FASE 4: Análise Avançada ---")
    metricas = {"analise_conteudo": {"in": 0, "out": 0}, "analise_global": {"in": 0, "out": 0}}
    fatos_e_descricoes = []
    inconsistencias_locais = []

    # --- SUB-FASE 4.1: VALIDAÇÕES LOCAIS (REGRAS) ---
    print("  FASE 4.1: Executando validações locais (baseadas em regras)...")
    for item in itens:
        if item["tipo"] == "texto":
            erros = validar_paragrafo_textual(item.get("texto_corrigido", item["texto"]))
            for erro in erros:
                erro['origem'] = item['origem']
                inconsistencias_locais.append(erro)
    for t_idx, table in enumerate(doc.tables):
        linhas_tabela = [[cell.text for cell in row.cells] for row in table.rows]
        erros = validar_tabela_diferenca(linhas_tabela)
        for erro in erros:
            erro['origem'] = f'Tabela {t_idx+1}'
            inconsistencias_locais.append(erro)
    print(f"    Validações locais concluídas. Encontrado(s) {len(inconsistencias_locais)} problema(s).")
    for inconsistencia in inconsistencias_locais:
        fatos_e_descricoes.append(f"- Origem ({inconsistencia['origem']}): [VALIDAÇÃO AUTOMÁTICA] {inconsistencia['detalhe']}")

    # --- SUB-FASE 4.2: ANÁLISE CRÍTICA E MULTIMODAL (IA) ---
    print("  FASE 4.2: Executando análise crítica e multimodal com IA...")
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        future_map = {}
        for item in itens:
            if item["tipo"] == "texto":
                future = executor.submit(executar_chamada_api, "", PROMPT_ANALISE_CRITICA_LOCAL.format(texto=item.get("texto_corrigido", item["texto"])))
                future_map[future] = item['origem']
            elif item["tipo"] == "imagem":
                future = executor.submit(analisar_imagem_com_visao, item["blob"], item["legenda_texto"])
                future_map[future] = item['origem']
        
        progresso = tqdm(as_completed(future_map), total=len(future_map), desc="Analisando conteúdo com IA", disable=DISABLE_TQDM)
        for future in progresso:
            origem = future_map[future]
            resultado, ti, to = future.result()
            metricas["analise_conteudo"]["in"] += ti
            metricas["analise_conteudo"]["out"] += to
            if resultado and resultado not in ["N/A"]:
                fatos_e_descricoes.append(f"- Origem ({origem}): {resultado}")
    
    # --- SUB-FASE 4.3: ANÁLISE GLOBAL FINAL (IA) ---
    if not fatos_e_descricoes:
        return "Nenhum fato relevante para análise global.", metricas, inconsistencias_locais
    
    dossie_final = "\n".join(fatos_e_descricoes)
    print("  FASE 4.3: Enviando dossiê consolidado para análise final...")
    relatorio_final, ti, to = executar_chamada_api("", PROMPT_ANALISE_GLOBAL_DOSSIE.format(dossie=dossie_final))
    metricas["analise_global"]["in"] += ti
    metricas["analise_global"]["out"] += to

    return relatorio_final, metricas, inconsistencias_locais

# --- Orquestração Principal ---
def processar_documento(nome_arquivo: str, id_revisao: int):
    caminho_doc = PASTA_ENTRADA / nome_arquivo
    nome_base = caminho_doc.stem
    pasta_saida_doc = PASTA_SAIDA / nome_base
    pasta_saida_doc.mkdir(parents=True, exist_ok=True)
    print(f"\n Iniciando processamento completo para: {nome_arquivo}")
    QueueEntry.objects.filter(id=id_revisao).update(progress=10)
    try:
        doc = Document(caminho_doc)
        itens_documento = mapear_conteudo_do_documento(doc)
        QueueEntry.objects.filter(id=id_revisao).update(progress=15)
        itens_categorizados = fase_1_mapeamento_e_categorizacao(itens_documento)
        QueueEntry.objects.filter(id=id_revisao).update(progress=30)
        itens_revisados, metricas, log_revisoes = fase_2_revisao_e_correcao(itens_categorizados)
        QueueEntry.objects.filter(id=id_revisao).update(progress=60)
        log_cruzamento = fase_3_verificacao_cruzada(itens_revisados)
        QueueEntry.objects.filter(id=id_revisao).update(progress=75)
        log_global, metricas_fase4, log_regras = fase_4_analise_multimodal_e_global(itens_revisados, doc)
        metricas.update(metricas_fase4)
        QueueEntry.objects.filter(id=id_revisao).update(progress=90)
        caminho_final = pasta_saida_doc / f"{nome_base}_revisao_completa.docx"
        doc.save(caminho_final)
        print(f" Documento revisado salvo em: {caminho_final}")
        plan_path = pasta_saida_doc / "avaliacao_completa.xlsx"
        wb = Workbook(); wb.remove(wb.active)
        ws_rev = wb.create_sheet("Log de Revisoes"); ws_rev.append(["Tipo", "Original", "Corrigido", "Validação Externa", "Justificativa"])
        for r in log_revisoes: ws_rev.append([r["tipo"], r["original"], r["corrigido"], r.get("validacao", ""), r.get("justificativa", "")])
        
        ws_regras = wb.create_sheet("Validações (Regras)")
        ws_regras.append(["Origem", "Regra", "Trecho Afetado", "Detalhe do Erro", "Gravidade"])
        for item in log_regras:
            ws_regras.append([item.get('origem'), item.get('regra'), item.get('trecho'), item.get('detalhe'), item.get('gravidade')])

        ws_cruz = wb.create_sheet("Verificação Cruzada"); ws_cruz.append(["Inconsistência de Citação Encontrada"])
        for i in log_cruzamento: ws_cruz.append([i])

        ws_glob = wb.create_sheet("Análise Global"); ws_glob.append(["Relatório de Inconsistências Globais"])
        ws_glob.append([log_global or "Nenhuma inconsistência global foi encontrada."])

        ws_custos = wb.create_sheet("Resumo_Custos"); ws_custos.append(["Processo", "Tokens In", "Tokens Out", "Custo USD", "Custo BRL"])
        # Loop para detalhar o custo de cada fase
        for fase, dados in metricas.items():
            if dados["in"] > 0 or dados["out"] > 0: # Apenas mostra fases que tiveram consumo
                usd = (dados["in"] * VALOR_INPUT + dados["out"] * VALOR_OUTPUT) / 1000
                nome_fase = fase.replace("_", " ").capitalize()
                ws_custos.append([f"Fase - {nome_fase}", dados["in"], dados["out"], round(usd, 4), round(usd * COTACAO_DOLAR, 2)])
        
        wb.save(plan_path)
        print(f" Planilha de resultados detalhada salva em: {plan_path}")
        
        relatorio_path = pasta_saida_doc / f"relatorio_tecnico_{nome_base}.docx"
        doc_relatorio = Document(); doc_relatorio.add_heading(f"Relatório Técnico de Revisão - {nome_base}", level=1)
        if log_revisoes:
            doc_relatorio.add_heading("Detalhes das Revisões", level=2)
            for i, r in enumerate(log_revisoes):
                p = doc_relatorio.add_paragraph(); p.add_run(f"Revisão {i+1} ({r['tipo']}): ").bold = True; p.add_run(r.get("justificativa", "Correção aplicada."))
                doc_relatorio.add_paragraph(f"Original: {r['original']}", style='Intense Quote'); doc_relatorio.add_paragraph(f"Corrigido: {r['corrigido']}", style='Intense Quote')
        if log_regras:
            doc_relatorio.add_heading("Validações Automáticas (Baseadas em Regras)", level=2)
            for item in log_regras:
                p = doc_relatorio.add_paragraph()
                p.add_run(f"Origem: {item.get('origem')} (Regra: {item.get('regra')})").bold = True
                p.add_run(f"\nDetalhe: {item.get('detalhe')}")
        if log_cruzamento:
            doc_relatorio.add_heading("Inconsistências de Citação", level=2)
            for i in log_cruzamento: doc_relatorio.add_paragraph(i, style='List Bullet')
        if log_global:
            doc_relatorio.add_heading("Análise de Contradições Globais e Multimodais", level=2); doc_relatorio.add_paragraph(log_global)
        doc_relatorio.save(relatorio_path)
        print(f" Relatório técnico salvo em: {relatorio_path}")
        return True
    except Exception:
        print(f" ERRO FATAL ao processar {nome_base}:"); traceback.print_exc()
        return False

if __name__ == "__main__":
    if not os.getenv("OPENAI_API_KEY"):
        print("ERRO: A variável de ambiente OPENAI_API_KEY não foi encontrada."); sys.exit(1)
    sucesso = processar_documento(args.arquivo, args.id_revisao)
    if not sucesso: sys.exit(1)
    print("\n Processamento completo finalizado com sucesso.")