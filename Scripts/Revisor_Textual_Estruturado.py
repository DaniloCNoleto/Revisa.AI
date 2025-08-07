# Scripts/Revisor_Textual_Estruturado.py
# Este script executa uma revisão textual simples e autocontida.
# Ele realiza seu próprio mapeamento e corrige o documento in-place.

import os
import sys
import time
import openai
import tiktoken
import re
import difflib
import traceback
from docx import Document
from openpyxl import Workbook
from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from dotenv import load_dotenv

# --- Bloco de Inicialização Django ---
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.append(str(PROJECT_ROOT))
load_dotenv(PROJECT_ROOT / ".env")
import django
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "dossel_project.settings")
django.setup()
import argparse
from revisor.models import QueueEntry
from django.conf import settings

# --- Argumentos e Configurações Globais ---
parser = argparse.ArgumentParser(description="Revisor Textual Estruturado Dossel.")
parser.add_argument("arquivo", help="Nome do arquivo DOCX (com extensão).")
parser.add_argument("--id_revisao", type=int, required=True, help="ID da revisão.")
args = parser.parse_args()

client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
ENCODER = tiktoken.encoding_for_model("gpt-4o")
MAX_WORKERS = min(10, (os.cpu_count() or 1) + 4)
TIMEOUT_API = 90
MAX_RETRIES = 2

PASTA_ENTRADA = Path(settings.PASTA_ENTRADA)
PASTA_SAIDA = Path(settings.PASTA_SAIDA)

# Desativa o TQDM se a saída não for um terminal interativo (ex: log de arquivo)
DISABLE_TQDM = not sys.stdout.isatty()

VALOR_INPUT = 0.005
VALOR_OUTPUT = 0.015
COTACAO_DOLAR = 5.65

# --- PROMPTS APRIMORADOS PARA MAIOR RIGOR ---
PROMPT_CATEGORIZACAO_SIMPLES = "Analise o parágrafo a seguir. Se ele contiver erros de gramática, ortografia, pontuação ou concordância, responda APENAS com a palavra 'textual'. Caso contrário, responda APENAS com 'nenhum'."
PROMPT_REVISAO_TEXTUAL = "Você é um revisor de textos acadêmicos e técnicos, especialista na norma culta da língua portuguesa. Sua tarefa é corrigir o trecho abaixo com o máximo rigor. Foco em erros de gramática, ortografia, concordância verbal e nominal, regência, crase e pontuação. NÃO altere o estilo de escrita, a voz do autor ou a terminologia técnica. Se o texto já estiver correto, repita-o integralmente. Responda APENAS com o texto corrigido."

# --- Funções Auxiliares ---
def contar_tokens(txt: str) -> int: return len(ENCODER.encode(txt))
def similaridade(a: str, b: str) -> float: return difflib.SequenceMatcher(None, a.lower(), b.lower()).ratio()

def executar_chamada_api(prompt_sistema: str, trecho_usuario: str):
    tokens_in, tokens_out = 0, 0
    for _ in range(MAX_RETRIES):
        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "system", "content": prompt_sistema}, {"role": "user", "content": trecho_usuario}],
                temperature=0.1, timeout=TIMEOUT_API
            )
            resposta_texto = response.choices[0].message.content.strip()
            tokens_in = response.usage.prompt_tokens
            tokens_out = response.usage.completion_tokens
            return resposta_texto, tokens_in, tokens_out
        except Exception:
            time.sleep(2)
    return None, tokens_in, tokens_out

# --- Funções do Pipeline ---
def coletar_blocos_de_texto(doc: Document):
    blocos = []
    def adicionar_paragrafos(container):
        for p in container:
            if p.text.strip():
                blocos.append({"objeto": p, "texto_original": p.text})
    adicionar_paragrafos(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                adicionar_paragrafos(cell.paragraphs)
    for section in doc.sections:
        adicionar_paragrafos(section.header.paragraphs)
        adicionar_paragrafos(section.footer.paragraphs)
    print(f"  FASE 0: Coleta concluída. {len(blocos)} blocos de texto encontrados.")
    return blocos

def revisar_documento(blocos: list):
    revisoes_log = []
    metricas = {"categorizacao": {"in": 0, "out": 0}, "revisao": {"in": 0, "out": 0}}

    # Fase 1: Mapeamento
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        future_map = {executor.submit(executar_chamada_api, "", PROMPT_CATEGORIZACAO_SIMPLES.format(texto=bloco["texto_original"])): bloco for bloco in blocos if len(bloco["texto_original"]) > 15}

        progresso = tqdm(as_completed(future_map), total=len(future_map), desc="FASE 1: Mapeando erros textuais", disable=DISABLE_TQDM)
        blocos_para_revisar = []
        for future in progresso:
            bloco_original = future_map[future]
            categoria, ti, to = future.result()
            metricas["categorizacao"]["in"] += ti
            metricas["categorizacao"]["out"] += to
            if categoria and 'textual' in categoria:
                blocos_para_revisar.append(bloco_original)

    # Fase 2: Correção
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        future_map_rev = {executor.submit(executar_chamada_api, PROMPT_REVISAO_TEXTUAL, bloco["texto_original"]): bloco for bloco in blocos_para_revisar}

        progresso_rev = tqdm(as_completed(future_map_rev), total=len(future_map_rev), desc="FASE 2: Corrigindo textos", disable=DISABLE_TQDM)
        for future in progresso_rev:
            bloco_original = future_map_rev[future]
            texto_corrigido, ti, to = future.result()
            metricas["revisao"]["in"] += ti
            metricas["revisao"]["out"] += to
            if texto_corrigido and similaridade(bloco_original["texto_original"], texto_corrigido) < 0.999:
                bloco_original["objeto"].text = texto_corrigido
                revisoes_log.append({"original": bloco_original["texto_original"], "corrigido": texto_corrigido})

    return metricas, revisoes_log

# --- Orquestração Principal ---
def processar_revisao_simples(nome_arquivo: str, id_revisao: int):
    caminho_doc = PASTA_ENTRADA / nome_arquivo
    nome_base = caminho_doc.stem
    pasta_saida_doc = PASTA_SAIDA / nome_base
    pasta_saida_doc.mkdir(parents=True, exist_ok=True)
    print(f"\n Iniciando Revisão Textual Estruturada para: {nome_arquivo}")
    
    try:
        QueueEntry.objects.filter(id=id_revisao).update(progress=10)
        doc = Document(caminho_doc)
        blocos = coletar_blocos_de_texto(doc)
        QueueEntry.objects.filter(id=id_revisao).update(progress=25)
        
        metricas, revisoes_log = revisar_documento(blocos)
        
        QueueEntry.objects.filter(id=id_revisao).update(progress=90)
        caminho_final = pasta_saida_doc / f"{nome_base}_revisao_simples.docx"
        doc.save(caminho_final)
        print(f" Documento com revisão simples salvo em: {caminho_final}")
        
        # Geração de Relatório em Planilha
        plan_path = pasta_saida_doc / "relatorio_revisao_simples.xlsx"
        wb = Workbook(); wb.remove(wb.active)
        ws_log = wb.create_sheet("Log de Revisoes"); ws_log.append(["Original", "Corrigido"])
        for r in revisoes_log: ws_log.append([r["original"], r["corrigido"]])
        
        ws_custos = wb.create_sheet("Resumo_Custos"); ws_custos.append(["Processo", "Tokens In", "Tokens Out", "Custo USD", "Custo BRL"])
        total_in = sum(v["in"] for v in metricas.values())
        total_out = sum(v["out"] for v in metricas.values())
        usd = (total_in * VALOR_INPUT + total_out * VALOR_OUTPUT) / 1000
        ws_custos.append(["Revisão Simples", total_in, total_out, round(usd, 4), round(usd * COTACAO_DOLAR, 2)])
        wb.save(plan_path)
        print(f" Planilha de resultados simples salva em: {plan_path}")

        # Geração do Relatório Técnico em .docx
        try:
            relatorio_path = pasta_saida_doc / f"relatorio_tecnico_simples_{nome_base}.docx"
            doc_relatorio = Document()
            doc_relatorio.add_heading(f"Relatório Técnico de Revisão Simples - {nome_base}", level=1)
            doc_relatorio.add_paragraph(f"Total de correções aplicadas: {len(revisoes_log)}")
            
            if revisoes_log:
                doc_relatorio.add_heading("Detalhes das Correções Aplicadas", level=2)
                for i, r in enumerate(revisoes_log):
                    p = doc_relatorio.add_paragraph()
                    p.add_run(f"Correção {i+1}:").bold = True
                    doc_relatorio.add_paragraph(f"Original: {r['original']}", style='Intense Quote')
                    doc_relatorio.add_paragraph(f"Corrigido: {r['corrigido']}", style='Intense Quote')
            
            doc_relatorio.save(relatorio_path)
            print(f" Relatório técnico (.docx) salvo em: {relatorio_path}")
        except Exception as e:
            print(f"  Aviso: Não foi possível gerar o relatório técnico em .docx. Erro: {e}")

        return True

    except Exception as e:
        print(f" ERRO FATAL ao processar {nome_base}:")
        traceback.print_exc()
        return False

if __name__ == "__main__":
    if not os.getenv("OPENAI_API_KEY"):
        print("ERRO: A variável de ambiente OPENAI_API_KEY não foi encontrada."); sys.exit(1)
    
    sucesso = processar_revisao_simples(args.arquivo, args.id_revisao)
    if not sucesso:
        sys.exit(1)
        
    print("\n Processo de revisão estruturada finalizado.")