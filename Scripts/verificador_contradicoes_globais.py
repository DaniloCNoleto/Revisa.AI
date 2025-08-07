# Scripts/verificador_contradicoes_globais.py
# VERSÃO UNIFICADA: Contém análise multimodal (texto e imagem) E validações locais baseadas em regras.

import os
import sys
import time
import openai
import tiktoken
import re
import base64
import traceback
from datetime import datetime
from typing import List, Dict, Any, Tuple, Optional
from docx import Document
from openpyxl import Workbook
from pathlib import Path
from dotenv import load_dotenv
from concurrent.futures import ThreadPoolExecutor, as_completed
from tqdm import tqdm

# --- Bloco de Inicialização Django ---
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.append(str(PROJECT_ROOT))
load_dotenv(PROJECT_ROOT / ".env")
import django
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "dossel_project.settings")
django.setup()
import argparse
from revisor.models import QueueEntry
from django.conf import settings

# --- Argumentos e Configurações Globais ---
parser = argparse.ArgumentParser(description="Verificador de Contradições Globais Dossel.")
parser.add_argument("arquivo", help="Nome do arquivo DOCX (com extensão).")
parser.add_argument("--id_revisao", type=int, required=True, help="ID da revisão.")
args = parser.parse_args()

client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
ENCODER = tiktoken.encoding_for_model("gpt-4o")
MAX_WORKERS = 5
TIMEOUT_API = 300
MAX_RETRIES = 2

# Desativa o TQDM se a saída não for um terminal interativo (ex: log de arquivo)
DISABLE_TQDM = not sys.stdout.isatty()

PASTA_ENTRADA = Path(settings.PASTA_ENTRADA)
PASTA_SAIDA = Path(settings.PASTA_SAIDA)

VALOR_INPUT = 0.005
VALOR_OUTPUT = 0.015
COTACAO_DOLAR = 5.65

# --- Prompts de Análise (Mantidos da sua versão com IA) ---
PROMPT_ANALISE_CRITICA_LOCAL = "Sua tarefa é atuar como um analista crítico. Leia o trecho a seguir e extraia as alegações factuais e dados chave. Além disso, se você identificar QUALQUER inconsistência lógica interna, contradição óbvia ou dado impossível (ex: '101% de satisfação'), aponte-a explicitamente. Seja conciso. Se não houver fatos ou inconsistências, responda com 'N/A'.\n\nTrecho: \"{texto}\""
PROMPT_DESCREVER_IMAGEM = "Sua tarefa é fornecer uma descrição objetiva e concisa dos principais elementos na imagem. Use o texto da legenda fornecida apenas como contexto. NÃO inclua nenhum prefixo como 'Figura X:'. Responda APENAS com a descrição do conteúdo visual.\n\nLegenda para Contexto: '{legenda}'"
PROMPT_ANALISE_GLOBAL_DOSSIE_V2 = "Você é um analista lógico. O dossiê abaixo contém uma lista de fatos, dados e descrições extraídos de um documento, incluindo validações automáticas. Sua tarefa é encontrar e listar TODAS as contradições lógicas ou factuais entre os itens desta lista. Se, e somente se, a lista de contradições estiver vazia, responda com a frase 'Nenhuma inconsistência global foi encontrada.'.\n\nDossiê para Análise:\n\n{dossie}"


################################################################################
### FUNÇÕES DE VALIDAÇÃO LOCAL (INCORPORADAS) ###
################################################################################

# ---------- utilitários ----------
def _br2float(valor: str) -> Optional[float]:
    if valor is None: return None
    s = str(valor).strip().replace(".", "").replace(" ", "").replace(",", ".")
    s = re.sub(r"[^0-9\.\-]", "", s)
    if s in ("", "-", ".", "-."): return None
    try: return float(s)
    except ValueError: return None

def _parse_hhmm(txt: str) -> Optional[Tuple[int, int]]:
    m = re.search(r"\b(\d{1,2}):(\d{2})\b", txt)
    if not m: return None
    h, mm = int(m.group(1)), int(m.group(2))
    if 0 <= h <= 23 and 0 <= mm <= 59: return h, mm
    return None

def _parse_date_br(txt: str) -> Optional[datetime]:
    m = re.search(r"\b(\d{2})/(\d{2})/(\d{4})\b", txt)
    if not m: return None
    d, mth, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
    try: return datetime(y, mth, d)
    except ValueError: return None

# ---------- Validações de Texto ----------
def validar_percentuais_impossiveis(texto: str) -> List[Dict[str, Any]]:
    achados = []
    for m in re.finditer(r"(\d{1,3}(?:[\.,]\d+)?)\s*%", texto):
        num = m.group(1).replace(".", "").replace(",", ".")
        try: val = float(num)
        except ValueError: continue
        if val > 100.0:
            achados.append({"regra": "percentual_impossivel", "trecho": m.group(0), "detalhe": f"Percentual {val}% > 100%", "gravidade": "alta"})
    return achados

def validar_periodo_do_dia(texto: str) -> List[Dict[str, Any]]:
    achados = []
    t, hm = texto.lower(), _parse_hhmm(t)
    if not hm: return achados
    h, _ = hm
    if "da manhã" in t and h >= 12: achados.append({"regra": "periodo_incompativel", "trecho": re.search(r"\b\d{1,2}:\d{2}\b.*?(manhã)", texto, flags=re.I).group(0), "detalhe": "Horário >= 12h marcado como 'da manhã'", "gravidade": "média"})
    if ("da tarde" in t or "da noite" in t) and h < 12: achados.append({"regra": "periodo_incompativel", "trecho": re.search(r"\b\d{1,2}:\d{2}\b.*?(tarde|noite)", texto, flags=re.I).group(0), "detalhe": "Horário < 12h marcado como 'da tarde/noite'", "gravidade": "média"})
    return achados

def validar_depois_antes_com_datas(texto: str) -> List[Dict[str, Any]]:
    achados = []
    datas = [m.group(0) for m in re.finditer(r"\b\d{2}/\d{2}/\d{4}\b", texto)]
    if len(datas) < 2: return achados
    d1, d2 = _parse_date_br(datas[0]), _parse_date_br(datas[1])
    if not d1 or not d2: return achados
    low = texto.lower()
    if "depois" in low and d1 <= d2: achados.append({"regra": "cronologia_incompativel", "trecho": texto.strip(), "detalhe": f"'{datas[0]}' não ocorre depois de '{datas[1]}'", "gravidade": "alta"})
    if "antes" in low and d1 >= d2: achados.append({"regra": "cronologia_incompativel", "trecho": texto.strip(), "detalhe": f"'{datas[0]}' não ocorre antes de '{datas[1]}'", "gravidade": "alta"})
    return achados

def validar_paragrafo_textual(texto: str) -> List[Dict[str, Any]]:
    issues = []
    issues.extend(validar_percentuais_impossiveis(texto))
    issues.extend(validar_periodo_do_dia(texto))
    issues.extend(validar_depois_antes_com_datas(texto))
    return issues

# ---------- Validação de Tabela ----------
HEADER_MAP = {"planejado": ["planejado", "orcado", "orçado"], "executado": ["executado", "realizado"], "diferenca": ["diferença", "diferenca", "variação", "variacao"], "total": ["total"]}

def _normaliza_header(h: str) -> str:
    hlow = h.strip().lower()
    for key, alts in HEADER_MAP.items():
        if any(a in hlow for a in alts): return key
    return hlow

def validar_tabela_diferenca(linhas_tabela: List[List[str]]) -> List[Dict[str, Any]]:
    inconsist = []
    if not linhas_tabela or len(linhas_tabela[0]) < 3: return inconsist
    headers = [_normaliza_header(h) for h in linhas_tabela[0]]
    try:
        i_plan, i_exec, i_dif = headers.index("planejado"), headers.index("executado"), headers.index("diferenca")
    except ValueError: return inconsist
    soma_plan = soma_exec = soma_dif = 0.0
    tem_total, total_row_idx = False, None
    for r_idx, row in enumerate(linhas_tabela[1:], start=1):
        if not row or all(not (c or "").strip() for c in row): continue
        if any((_normaliza_header(c) == "total") for c in row):
            tem_total, total_row_idx = True, r_idx
            continue
        v_plan = _br2float(row[i_plan]) if i_plan < len(row) else None
        v_exec = _br2float(row[i_exec]) if i_exec < len(row) else None
        v_dif  = _br2float(row[i_dif])  if i_dif  < len(row) else None
        if v_plan is not None: soma_plan += v_plan
        if v_exec is not None: soma_exec += v_exec
        if v_dif  is not None: soma_dif  += v_dif
        if None not in (v_plan, v_exec, v_dif):
            calc = v_exec - v_plan
            if abs(calc - v_dif) > 0.009:
                inconsist.append({"regra": "diferenca_incorreta", "trecho": f"linha {r_idx+1}", "detalhe": f"Executado({v_exec}) - Planejado({v_plan}) = {calc:.2f}, dif. informada = {v_dif:.2f}", "gravidade": "alta"})
    if tem_total and total_row_idx is not None:
        row = linhas_tabela[total_row_idx]
        v_plan_t, v_exec_t, v_dif_t = (_br2float(row[i_plan]) if i_plan < len(row) else None, _br2float(row[i_exec]) if i_exec < len(row) else None, _br2float(row[i_dif])  if i_dif  < len(row) else None)
        if v_plan_t is not None and abs(v_plan_t - soma_plan) > 0.009: inconsist.append({"regra": "total_planejado_incorreto", "trecho": f"linha {total_row_idx+1}", "detalhe": f"TOTAL Planejado {v_plan_t:.2f} ≠ soma {soma_plan:.2f}", "gravidade": "alta"})
        if v_exec_t is not None and abs(v_exec_t - soma_exec) > 0.009: inconsist.append({"regra": "total_executado_incorreto", "trecho": f"linha {total_row_idx+1}", "detalhe": f"TOTAL Executado {v_exec_t:.2f} ≠ soma {soma_exec:.2f}", "gravidade": "alta"})
        if v_dif_t is not None and abs(v_dif_t - soma_dif) > 0.009: inconsist.append({"regra": "total_diferenca_incorreto", "trecho": f"linha {total_row_idx+1}", "detalhe": f"TOTAL Diferença {v_dif_t:.2f} ≠ soma {soma_dif:.2f}", "gravidade": "alta"})
    return inconsist

################################################################################
### FUNÇÕES DE PIPELINE (ATUALIZADAS) ###
################################################################################

def executar_chamada_api(prompt_sistema: str, trecho_usuario: str):
    # (Mantida da sua versão A)
    tokens_in, tokens_out = 0, 0
    for _ in range(MAX_RETRIES):
        try:
            response = client.chat.completions.create(model="gpt-4o", messages=[{"role": "system", "content": prompt_sistema}, {"role": "user", "content": trecho_usuario}], temperature=0.1, timeout=TIMEOUT_API)
            resposta_texto = response.choices[0].message.content.strip()
            tokens_in, tokens_out = response.usage.prompt_tokens, response.usage.completion_tokens
            return resposta_texto, tokens_in, tokens_out
        except Exception: time.sleep(2)
    return "Falha na chamada à API.", 0, 0

def analisar_imagem_com_visao(blob_imagem, legenda):
    # (Mantida da sua versão A)
    try:
        prompt_texto = PROMPT_DESCREVER_IMAGEM.format(legenda=legenda or "Nenhuma")
        response = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": [{"type": "text", "text": prompt_texto}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64.b64encode(blob_imagem).decode('utf-8')}"}}]}], max_tokens=300)
        tokens_in, tokens_out = response.usage.prompt_tokens, response.usage.completion_tokens
        return response.choices[0].message.content, tokens_in, tokens_out
    except Exception as e:
        print(f"  Aviso: Falha ao analisar imagem. Erro: {e}")
        return "Falha ao analisar a imagem.", 0, 0

def mapear_conteudo_do_documento(doc: Document):
    # (Mantida e corrigida da sua versão A)
    print("  Mapeando conteúdo (parágrafos, tabelas, imagens)...")
    itens_mapeados = []
    regex_legenda = re.compile(r"^(figura|fig|foto|imagem)[\s-]?\d+.*", re.IGNORECASE)
    
    # Mapeia parágrafos e imagens
    i = 0
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        texto_paragrafo = p.text.strip()
        imagens_blob = [part.blob for run in p.runs for inline in run.element.xpath('.//wp:inline') for rId in inline.xpath('.//a:blip/@r:embed') if (part := doc.part.related_parts.get(rId))]
        if imagens_blob:
            legenda_texto = doc.paragraphs[i+1].text.strip() if i + 1 < len(doc.paragraphs) and regex_legenda.match(doc.paragraphs[i+1].text.strip()) else ""
            if legenda_texto: i += 1
            for blob in imagens_blob:
                itens_mapeados.append({"tipo": "imagem", "blob": blob, "legenda_texto": legenda_texto, "origem": legenda_texto or f"Imagem P{len(itens_mapeados)+1}"})
        elif texto_paragrafo:
            itens_mapeados.append({"tipo": "texto", "texto": texto_paragrafo, "origem": f"Parágrafo {len(itens_mapeados)+1}"})
        i += 1
        
    # Mapeia tabelas (como texto e como objeto)
    for t_idx, table in enumerate(doc.tables):
        itens_mapeados.append({"tipo": "tabela", "objeto_tabela": table, "origem": f"Tabela {t_idx+1}"})
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    itens_mapeados.append({'tipo': 'texto', 'texto': cell.text.strip(), 'origem': f'Tabela {t_idx+1}, L{r_idx+1}, C{c_idx+1}'})

    print(f"  Mapeamento concluído. {len(itens_mapeados)} itens de conteúdo etiquetados.")
    return itens_mapeados


################################################################################
### FUNÇÃO PRINCIPAL DO SCRIPT (UNIFICADA) ###
################################################################################

def analisar_documento(nome_arquivo: str, id_revisao: int):
    caminho_doc = PASTA_ENTRADA / nome_arquivo
    nome_base = caminho_doc.stem
    pasta_saida_doc = PASTA_SAIDA / nome_base
    pasta_saida_doc.mkdir(parents=True, exist_ok=True)
    print(f"\n Iniciando Verificador de Contradições para: {nome_arquivo}")

    try:
        QueueEntry.objects.filter(id=id_revisao).update(progress=5)
        doc = Document(caminho_doc)
        itens_documento = mapear_conteudo_do_documento(doc)
        QueueEntry.objects.filter(id=id_revisao).update(progress=15)
        
        # --- FASE 1: VALIDAÇÕES LOCAIS (REGRAS) ---
        print("\n--- FASE 1: Executando validações locais (baseadas em regras) ---")
        inconsistencias_locais = []
        for item in itens_documento:
            if item["tipo"] == "texto":
                erros = validar_paragrafo_textual(item["texto"])
                for erro in erros:
                    erro['origem'] = item['origem']
                    inconsistencias_locais.append(erro)
            elif item["tipo"] == "tabela":
                linhas_tabela = [[cell.text for cell in row.cells] for row in item["objeto_tabela"].rows]
                erros = validar_tabela_diferenca(linhas_tabela)
                for erro in erros:
                    erro['origem'] = item['origem']
                    inconsistencias_locais.append(erro)
        print(f"  Validações locais concluídas. Encontrado(s) {len(inconsistencias_locais)} problema(s).")
        QueueEntry.objects.filter(id=id_revisao).update(progress=30)
        
        # --- FASE 2: ANÁLISE CRÍTICA E MULTIMODAL (IA) ---
        print("\n--- FASE 2: Executando análise crítica e multimodal (IA) ---")
        metricas = {"extracao_ia": {"in": 0, "out": 0}, "analise_final": {"in": 0, "out": 0}}
        fatos_e_descricoes = []
        
        # Adiciona os resultados das regras ao dossiê
        for inconsistencia in inconsistencias_locais:
            fatos_e_descricoes.append(f"- Origem ({inconsistencia['origem']}): [VALIDAÇÃO AUTOMÁTICA] {inconsistencia['detalhe']}")

        # Executa a análise da IA em paralelo
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            future_map = {}
            for item in itens_documento:
                if item["tipo"] == "texto" and len(item["texto"]) > 15:
                    future = executor.submit(executar_chamada_api, "", PROMPT_ANALISE_CRITICA_LOCAL.format(texto=item["texto"]))
                    future_map[future] = item['origem']
                elif item["tipo"] == "imagem":
                    future = executor.submit(analisar_imagem_com_visao, item["blob"], item["legenda_texto"])
                    future_map[future] = item['origem']

            progresso = tqdm(as_completed(future_map), total=len(future_map), desc="Analisando conteúdo com IA", disable=DISABLE_TQDM)
            for future in progresso:
                origem = future_map[future]
                resultado, ti, to = future.result()
                metricas["extracao_ia"]["in"] += ti
                metricas["extracao_ia"]["out"] += to
                if resultado and resultado not in ["N/A"]:
                    fatos_e_descricoes.append(f"- Origem ({origem}): {resultado}")

        QueueEntry.objects.filter(id=id_revisao).update(progress=75)

        # --- FASE 3: ANÁLISE GLOBAL FINAL (IA) ---
        print("\n--- FASE 3: Consolidando e enviando dossiê para análise final ---")
        if not fatos_e_descricoes:
            relatorio_global = "Nenhum fato, imagem ou erro relevante para análise global."
        else:
            dossie_final = "\n".join(fatos_e_descricoes)
            relatorio_global, ti, to = executar_chamada_api("", PROMPT_ANALISE_GLOBAL_DOSSIE_V2.format(dossie=dossie_final))
            metricas["analise_final"]["in"] += ti
            metricas["analise_final"]["out"] += to
        QueueEntry.objects.filter(id=id_revisao).update(progress=90)
        
        # --- FASE 4: GERAÇÃO DE RELATÓRIOS ---
        print("\n--- FASE 4: Gerando relatórios consolidados ---")
        plan_path = pasta_saida_doc / f"relatorio_inconsistencias_{nome_base}.xlsx"
        wb = Workbook()
        
        # Aba 1: Validações Locais (Regras)
        ws_locais = wb.active
        ws_locais.title = "Validações Locais (Regras)"
        ws_locais.append(["Origem", "Regra", "Trecho Afetado", "Detalhe do Erro", "Gravidade"])
        if inconsistencias_locais:
            for item in inconsistencias_locais:
                ws_locais.append([item.get('origem'), item.get('regra'), item.get('trecho'), item.get('detalhe'), item.get('gravidade')])
        else:
            ws_locais.append(["Nenhuma inconsistência encontrada pelas regras automáticas."])

        # Aba 2: Análise Global (IA)
        ws_global = wb.create_sheet("Análise Global (IA)")
        ws_global.append(["Relatório de Contradições Lógicas Encontradas pela IA no Dossiê"])
        ws_global.append([relatorio_global or "Nenhuma inconsistência global foi encontrada."])

        # Aba 3: Custos
        ws_custos = wb.create_sheet("Resumo_Custos")
        ws_custos.append(["Processo", "Tokens In", "Tokens Out", "Custo USD", "Custo BRL"])
        total_in = sum(v["in"] for v in metricas.values())
        total_out = sum(v["out"] for v in metricas.values())
        usd = (total_in * VALOR_INPUT + total_out * VALOR_OUTPUT) / 1000
        ws_custos.append(["Análise de Inconsistências Total", total_in, total_out, round(usd, 4), round(usd * COTACAO_DOLAR, 2)])
        
        wb.save(plan_path)
        print(f" Planilha de análise e custos salva em: {plan_path}")

        # Relatório Técnico .docx
        relatorio_path = pasta_saida_doc / f"relatorio_tecnico_inconsistencias_{nome_base}.docx"
        doc_relatorio = Document()
        doc_relatorio.add_heading(f"Relatório de Análise de Inconsistências - {nome_base}", level=1)
        if inconsistencias_locais:
            doc_relatorio.add_heading("1. Validações Automáticas (Baseadas em Regras)", level=2)
            for item in inconsistencias_locais:
                p = doc_relatorio.add_paragraph()
                p.add_run(f"Origem: {item.get('origem')} (Regra: {item.get('regra')})").bold = True
                p.add_run(f"\nDetalhe: {item.get('detalhe')}")
        doc_relatorio.add_heading("2. Análise Global de Contradições (IA)", level=2)
        doc_relatorio.add_paragraph(relatorio_global)
        doc_relatorio.save(relatorio_path)
        print(f" Relatório técnico salvo em: {relatorio_path}")

        return True
    except Exception:
        print(f" ERRO FATAL ao analisar {nome_base}:")
        traceback.print_exc()
        return False

if __name__ == "__main__":
    if not os.getenv("OPENAI_API_KEY"):
        print("ERRO: A variável de ambiente OPENAI_API_KEY não foi encontrada."); sys.exit(1)
    sucesso = analisar_documento(args.arquivo, args.id_revisao)
    if not sucesso:
        sys.exit(1)
    print("\n Processo de verificação de contradições finalizado.")