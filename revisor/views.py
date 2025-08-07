from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from django.contrib.auth.models import User
from django.urls import reverse
from django.conf import settings
from django.contrib.auth.decorators import login_required
import re
from httpcore import request
import pandas as pd
import plotly.express as px
import json
import plotly.graph_objects as go
from plotly.offline import plot

from .models import QueueEntry # Mantido

import os # Mantido
import shutil # Adicionado para operações de arquivos e pastas
import sys  # Adicionado para uso de sys.executable
import subprocess  # Adicionado para uso de subprocess.Popen
from pathlib import Path
from django.core.files.storage import FileSystemStorage
from .forms import UploadForm  # Removido porque forms.py não existe ou não é necessário
import subprocess

from django.core.mail import send_mail
from django.template.loader import render_to_string
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.utils.encoding import force_bytes, force_str
from django.contrib.sites.shortcuts import get_current_site
from django.contrib.auth.tokens import default_token_generator
from docx import Document



def contar_palavras_docx(caminho_arquivo):
    """Abre um arquivo .docx e conta o número total de palavras."""
    try:
        doc = Document(caminho_arquivo)
        total_palavras = 0
        for para in doc.paragraphs:
            total_palavras += len(para.text.split())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        total_palavras += len(para.text.split())
        return total_palavras
    except Exception:
        return 0

def parse_contradiction_report(report_text):
    """Quebra o relatório de texto da IA em uma lista de contradições estruturadas."""
    if not report_text or "nenhuma inconsistência" in report_text.lower():
        return []
    
    contradictions = re.split(r'-\s*Contradição sobre', report_text)
    parsed_list = []
    for contradiction in contradictions:
        if not contradiction.strip(): continue
        
        tema_match = re.match(r'(.+?):', contradiction, re.DOTALL)
        trecho1_match = re.search(r'Trecho 1:\s*"(.*?)"', contradiction, re.DOTALL)
        trecho2_match = re.search(r'Trecho 2:\s*"(.*?)"', contradiction, re.DOTALL)
        
        if tema_match and trecho1_match and trecho2_match:
            parsed_list.append({
                'tema': tema_match.group(1).strip(),
                'trecho1': trecho1_match.group(1).strip(),
                'trecho2': trecho2_match.group(1).strip(),
            })
    return parsed_list
    
# --- VIEWS DE AUTENTICAÇÃO ---
def login_view(request):
    """
    Handles user login.
    """
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')

        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
            next_url = request.GET.get('next')
            if next_url:
                return redirect(next_url)
            else:
                return redirect('upload')
        else:
            messages.error(request, 'Nome de usuário ou senha inválidos.')
            return render(request, 'revisor/login.html', {'error': 'Nome de usuário ou senha inválidos.'})
    else:
        return render(request, 'revisor/login.html')

def register_view(request):
    """
    Handles user registration with validation and email confirmation.
    """
    if request.method == 'POST':
        full_name = request.POST.get('full_name')
        username = request.POST.get('username')
        email = request.POST.get('email')
        password = request.POST.get('password')
        password_confirm = request.POST.get('password_confirm')

        errors = []

        if not all([full_name, username, email, password, password_confirm]):
            errors.append('Todos os campos são obrigatórios.')

        if User.objects.filter(username=username).exists():
            errors.append('Usuário já existe.')

        if User.objects.filter(email=email).exists():
            errors.append('E-mail já cadastrado.')

        if not email.endswith('@dosselambiental.com.br'):
            errors.append('Use seu e-mail "@dosselambiental.com.br".')

        if password != password_confirm:
            errors.append('As senhas não coincidem.')

        if len(password) < 8:
            errors.append('A senha deve ter no mínimo 8 caracteres.')
        if not re.search(r'[!@#$%^&*(),.?":{}|<>]', password):
            errors.append('A senha deve conter ao menos um caractere especial.')
        if not re.search(r'\d', password):
            errors.append('A senha deve conter ao menos um número.')
        if not re.search(r'[A-Z]', password):
            errors.append('A senha deve conter ao menos uma letra maiúscula.')

        if errors:
            for error_msg in errors:
                messages.error(request, error_msg)
            return render(request, 'revisor/register.html', {'error': 'Verifique os erros no formulário.'})

        try:
            user = User.objects.create_user(username=username, email=email, password=password)
            user.first_name = full_name
            user.is_active = False
            user.save()

            current_site = get_current_site(request)
            mail_subject = 'Ative sua conta na Dossel Ambiental'
            message = render_to_string('revisor/acc_active_email.html', {
                'user': user,
                'domain': current_site.domain,
                'uid': urlsafe_base64_encode(force_bytes(user.pk)),
                'token': default_token_generator.make_token(user),
            })
            send_mail(mail_subject, message, settings.DEFAULT_FROM_EMAIL, [user.email], fail_silently=False)

            messages.info(request, 'Por favor, confirme seu e-mail para completar o registro. Verifique sua caixa de entrada, incluindo spam.')
            return redirect('email_sent_confirmation')

        except Exception as e:
            messages.error(request, f"Erro inesperado ao registrar: {e}")
            return render(request, 'revisor/register.html', {'error': f"Erro inesperado ao registrar: {e}"})

    return render(request, 'revisor/register.html')


def activate(request, uidb64, token):
    try:
        uid = force_str(urlsafe_base64_decode(uidb64))
        user = User.objects.get(pk=uid)
    except (TypeError, ValueError, OverflowError, User.DoesNotExist):
        user = None

    if user is not None and default_token_generator.check_token(user, token):
        user.is_active = True
        user.save()
        messages.success(request, 'Sua conta foi ativada com sucesso! Você já pode fazer login.')
        return redirect('login')
    else:
        messages.error(request, 'Link de ativação inválido ou expirado! Por favor, tente novamente.')
        return render(request, 'revisor/activation_invalid.html', {'error': 'Link inválido ou expirado.'})


def email_sent_confirmation_view(request):
    return render(request, 'revisor/email_sent_confirmation.html')

def activation_invalid_view(request):
    return render(request, 'revisor/activation_invalid.html')


def logout_view(request):
    """
    Handles user logout.
    """
    logout(request)
    messages.info(request, 'Você foi desconectado(a).')
    return redirect('login')

# --- VIEWS PROTEGIDAS COM @login_required ---

@login_required
def upload_view(request):
    if request.method == 'POST' and request.FILES.get('documento'):
        documento = request.FILES['documento']

        # Define o caminho para guardar o ficheiro (ex: numa pasta 'uploads' dentro de MEDIA_ROOT)
        # Certifique-se que MEDIA_ROOT e MEDIA_URL estão configurados em settings.py
        upload_path = os.path.join(settings.MEDIA_ROOT, 'uploads')
        os.makedirs(upload_path, exist_ok=True) # Cria a pasta se não existir
        
        fs = FileSystemStorage(location=upload_path)
        filename = fs.save(documento.name, documento)
        file_path = fs.path(filename)

        # Guarda o caminho do ficheiro e o nome original na sessão do utilizador
        request.session['uploaded_file_path'] = file_path
        request.session['uploaded_file_name'] = filename # Use o nome salvo pelo FileSystemStorage para evitar colisões
        
        messages.success(request, "Documento enviado com sucesso! Selecione o modo de revisão.")
        return redirect('modo_revisao')

    return render(request, 'revisor/upload.html')


@login_required
def modo_revisao_view(request):
    file_path_session = request.session.get('uploaded_file_path')
    file_name = request.session.get('uploaded_file_name')

    if not file_path_session or not os.path.exists(file_path_session):
        messages.error(request, "Por favor, faça primeiro o upload de um documento válido.")
        return redirect('upload')

    if request.method == 'POST':
        modo_revisao = request.POST.get('modo') 
        if not modo_revisao:
            messages.error(request, "Por favor, selecione um modo de revisão.")
            return render(request, 'revisor/modo_revisao.html', {'file_name': file_name})

        caminho_final_entrada = os.path.join(settings.PASTA_ENTRADA, file_name)
        try:
            shutil.copy(file_path_session, caminho_final_entrada)
        except Exception as e:
            messages.error(request, f"Erro ao mover o arquivo para a área de processamento: {e}")
            return redirect('upload')

        nova_revisao = QueueEntry.objects.create(
            user=request.user, file_name=file_name,
            file_path=caminho_final_entrada,
            status='na_fila', mode=modo_revisao
        )

        script_path = os.path.join(settings.BASE_DIR, 'Scripts', 'Gerenciador_Revisores_Estruturados.py')
        
        # --- CORREÇÃO FINAL NA CONSTRUÇÃO DOS ARGUMENTOS ---
        # A flag e seu valor devem ser itens separados na lista.
        args = [
            sys.executable,
            '-u',
            script_path,
            file_name, # CORRETO: Argumento posicional para o nome do arquivo
            '--id_revisao', str(nova_revisao.id),
            '--email', request.user.email,
            '--modo', modo_revisao
        ]
        # --- FIM DA CORREÇÃO ---

        try:
            log_path = os.path.join(settings.BASE_DIR, 'logs')
            os.makedirs(log_path, exist_ok=True)
            stdout_log = open(os.path.join(log_path, f'revisao_{nova_revisao.id}_stdout.log'), 'w')
            stderr_log = open(os.path.join(log_path, f'revisao_{nova_revisao.id}_stderr.log'), 'w')

            subprocess.Popen(args, stdout=stdout_log, stderr=stderr_log, cwd=settings.BASE_DIR)
            
            nova_revisao.status = 'processando'
            nova_revisao.save()

            if 'uploaded_file_path' in request.session: del request.session['uploaded_file_path']
            if 'uploaded_file_name' in request.session: del request.session['uploaded_file_name']

            messages.info(request, f"A revisão de '{file_name}' foi iniciada.")
            return redirect('acompanhamento')
        except Exception as e:
            messages.error(request, f"Erro CRÍTICO ao tentar iniciar o subprocesso: {e}")
            nova_revisao.status = 'erro'
            nova_revisao.save()
            return redirect('upload')

    return render(request, 'revisor/modo_revisao.html', {'file_name': file_name})
@login_required
def acompanhamento_view(request):
    # Primeiro, procuramos por uma revisão que ainda esteja com o status 'processando'
    revisao_em_andamento = QueueEntry.objects.filter(status='processando', user=request.user).order_by('-created_at').first()

    if revisao_em_andamento:
        # Se encontrarmos uma, verificamos o progresso. Se já estiver em 100%, redirecionamos.
        # Isso cobre o caso em que o Gerenciador ainda não atualizou o status para 'concluido'.
        if revisao_em_andamento.progress >= 100:
            return redirect('resultados', revision_id=revisao_em_andamento.id)
        
        # Se não, mostramos a página de acompanhamento normal
        fila_de_espera = QueueEntry.objects.filter(status='na_fila', user=request.user).order_by('created_at')
        context = {
            'revisao_em_andamento': revisao_em_andamento,
            'fila_de_espera': fila_de_espera,
            'progresso': revisao_em_andamento.progress,
        }
        return render(request, 'revisor/acompanhamento.html', context)
    else:
        # Se NÃO houver nenhuma revisão com status 'processando', isso pode significar duas coisas:
        # 1. Não há nada na fila.
        # 2. Uma revisão ACABOU de ser concluída (status mudou para 'concluido').

        # Então, procuramos pela última revisão concluída do usuário.
        revisao_concluida_recente = QueueEntry.objects.filter(user=request.user, status='concluido').order_by('-created_at').first()        
        # Se encontrarmos uma revisão concluída e o usuário veio para esta página
        # sem um ID específico, o redirecionamos para o resultado mais recente.
        if revisao_concluida_recente:
            # Para evitar redirecionar para resultados muito antigos, podemos adicionar um
            # filtro de tempo, mas para a funcionalidade principal, o redireto é suficiente.
            return redirect('resultados', revision_id=revisao_concluida_recente.id)
            
        # Se não há nada processando e nenhuma concluída recentemente, apenas renderiza a página vazia.
        context = {
            'revisao_em_andamento': None,
            'fila_de_espera': [],
            'progresso': 0,
        }
        return render(request, 'revisor/acompanhamento.html', context)

# --- AJUSTE CRÍTICO: `resultados_view` agora usa `revision_id` ---


@login_required
def cancelar_revisao_view(request, revision_id):
    if request.method == 'POST':
        try:
            revisao_para_cancelar = QueueEntry.objects.get(id=revision_id, user=request.user)
            
            # Cria a pasta de sinais de cancelamento se não existir
            os.makedirs(settings.PASTA_CANCELAMENTO, exist_ok=True)
            
            # Cria um arquivo vazio que o script irá verificar. Ex: 'cancel_123.txt'
            sinal_cancelamento_path = os.path.join(settings.PASTA_CANCELAMENTO, f'cancel_{revision_id}.txt')
            with open(sinal_cancelamento_path, 'w') as f:
                f.write('cancel')
            
            # Atualiza o status no banco para refletir o cancelamento
            revisao_para_cancelar.status = 'cancelado'
            revisao_para_cancelar.save()
            
            messages.info(request, f'Sinal de cancelamento enviado para a revisão de "{revisao_para_cancelar.file_name}". O processo será interrompido em breve.')

        except QueueEntry.DoesNotExist:
            messages.error(request, 'Revisão não encontrada ou você não tem permissão para cancelá-la.')
        except Exception as e:
            messages.error(request, f'Ocorreu um erro ao tentar cancelar a revisão: {e}')
            
    return redirect('upload')


@login_required
def resultados_view(request, revision_id):
    try:
        revisao = QueueEntry.objects.get(id=revision_id, user=request.user)
    except QueueEntry.DoesNotExist:
        messages.error(request, "Revisão não encontrada ou não pertence a você.")
        return redirect('historico')

    # --- SETUP INICIAL ---
    context = {'revisao': revisao}
    nome_base = Path(revisao.file_name).stem
    pasta_saida = Path(settings.MEDIA_ROOT) / "saida" / nome_base

    # --- LÓGICA CONDICIONAL POR MODO DE REVISÃO ---
    
    # 1. MODO SIMPLES
    if revisao.mode == 'simples':
        try:
            planilha_path = pasta_saida / "relatorio_revisao_simples.xlsx"
            df_revisoes = pd.read_excel(planilha_path, sheet_name="Log de Revisoes")
            
            palavras_corrigidas = int(df_revisoes['Corrigido'].astype(str).apply(lambda x: len(x.split())).sum())
            palavras_totais = contar_palavras_docx(revisao.file_path)
            palavras_nao_corrigidas = max(0, palavras_totais - palavras_corrigidas)

            fig_donut = go.Figure(data=[go.Pie(labels=['Corrigidas', 'Não Corrigidas'], values=[palavras_corrigidas, palavras_nao_corrigidas], hole=.6, marker_colors=['#00af74', '#444444'])])
            fig_donut.update_traces(hoverinfo='label+percent', textinfo='value', textfont_size=16)
            fig_donut.update_layout(
                title_text='Proporção de Palavras Corrigidas', 
                paper_bgcolor='rgba(0,0,0,0)', 
                plot_bgcolor='rgba(0,0,0,0)', 
                font=dict(color='white'), 
                legend=dict(orientation="h", yanchor="bottom", y=-0.1, xanchor="center", x=0.5),
                height=400  # <<< ALTURA AJUSTADA
            )
            
            context['plot_div'] = fig_donut.to_html(full_html=False, include_plotlyjs=False)
            context['palavras_corrigidas'] = palavras_corrigidas
            context['palavras_entrada'] = palavras_totais

        except Exception as e:
            messages.error(request, f"Erro ao gerar o resultado da Revisão Simples: {e}")

    # 2. MODO INCONSISTÊNCIAS
    elif revisao.mode == 'inconsistencias':
        try:
            planilha_path = pasta_saida / f"relatorio_inconsistencias_{nome_base}.xlsx"
            df_analise = pd.read_excel(planilha_path, sheet_name="Análise Global")
            relatorio_texto = df_analise.iloc[0, 0] if not df_analise.empty else ""
            
            num_inconsistencias = len(parse_contradiction_report(relatorio_texto))

            context['kpi'] = {'valor': num_inconsistencias, 'label': 'Inconsistências Encontradas'}
            context['palavras_entrada'] = contar_palavras_docx(revisao.file_path)

        except Exception as e:
            messages.error(request, f"Erro ao gerar o resultado de Inconsistências: {e}")

    # 3. MODO COMPLETA
    elif revisao.mode == 'completa':
        try:
            planilha_path = pasta_saida / "avaliacao_completa.xlsx"
            df_revisoes = pd.read_excel(planilha_path, sheet_name="Log de Revisoes")
            contagem_por_tipo = df_revisoes['Tipo'].value_counts()

            # KPI para Inconsistências
            total_inconsistencias = contagem_por_tipo.get('Inconsistencia', 0)
            context['kpi_inconsistencias'] = {'valor': total_inconsistencias, 'label': 'Inconsistências Resolvidas'}

            # Gráfico de Barras para os outros tipos
            dados_outros = {
                'Tipo': ['Textual', 'Bibliografia', 'Estrutura'],
                'Total': [contagem_por_tipo.get('Textual', 0), contagem_por_tipo.get('Bibliografico', 0), contagem_por_tipo.get('Estrutura', 0)]
            }
            df_outros = pd.DataFrame(dados_outros).sort_values(by='Total', ascending=True)
            fig_outros = px.bar(df_outros, x='Total', y='Tipo', orientation='h', text='Total', title='Análise de Correções', color_discrete_sequence=['#00af74', '#5A4A2F', '#22a6b3'])
            fig_outros.update_layout(
                paper_bgcolor='rgba(0,0,0,0)', 
                plot_bgcolor='rgba(0,0,0,0)', 
                font=dict(color='white'), 
                yaxis=dict(title_text=''),
                height=400  # <<< ALTURA AJUSTADA
            )
            
            context['plot_correcoes'] = fig_outros.to_html(full_html=False, include_plotlyjs=False)
            context['palavras_corrigidas'] = int(df_revisoes.shape[0])
            context['palavras_entrada'] = contar_palavras_docx(revisao.file_path)

        except Exception as e:
            messages.error(request, f"Erro ao gerar o resultado da Revisão Completa: {e}")

    # --- LÓGICA DE DOWNLOADS (AJUSTADA CONFORME SOLICITADO) ---
    base_url_para_downloads = f'{settings.MEDIA_URL}saida/{nome_base}/'
    arquivos_para_download = []

    if revisao.mode == 'simples':
        arquivos_para_download.extend([
            {'label': 'Documento Revisado', 'file': f"{base_url_para_downloads}{nome_base}_revisao_simples.docx"},
            {'label': 'Relatório Técnico', 'file': f"{base_url_para_downloads}relatorio_tecnico_simples_{nome_base}.docx"}
        ])
    elif revisao.mode == 'completa':
        arquivos_para_download.extend([
            {'label': 'Documento Revisado', 'file': f"{base_url_para_downloads}{nome_base}_revisao_completa.docx"},
            {'label': 'Relatório Técnico', 'file': f"{base_url_para_downloads}relatorio_tecnico_{nome_base}.docx"}
        ])
    elif revisao.mode == 'inconsistencias':
        arquivos_para_download.extend([
            {'label': 'Relatório de Inconsistências', 'file': f"{base_url_para_downloads}relatorio_tecnico_inconsistencias_{nome_base}.docx"},
            {'label': 'Planilha de Análise', 'file': f"{base_url_para_downloads}relatorio_inconsistencias_{nome_base}.xlsx"}
        ])

    context['arquivos'] = arquivos_para_download
    
    return render(request, 'revisor/resultados.html', context)

# --- IMPLEMENTAÇÃO DA `historico_view` ---
@login_required
def historico_view(request):
    # Busca todas as revisões finalizadas (concluídas ou canceladas)
    revisoes_finalizadas = QueueEntry.objects.filter(
        user=request.user, 
        status__in=['concluido', 'cancelado']
    ).order_by('-created_at')
    
    historico_list = []
    for rev in revisoes_finalizadas:
        nome_base = Path(rev.file_name).stem
        base_url = f'{settings.MEDIA_URL}saida/{nome_base}/'
        
        # --- LÓGICA ADAPTATIVA PARA DOWNLOADS ---
        downloads = []
        if rev.status == 'concluido':
            if rev.mode == 'completa':
                downloads.append({'label': 'Doc Revisado (Completo)', 'file': f'{base_url}{nome_base}_revisao_completa.docx'})
                downloads.append({'label': 'Relatório Técnico', 'file': f'{base_url}relatorio_tecnico_{nome_base}.docx'})
                downloads.append({'label': 'Planilha de Análise', 'file': f'{base_url}avaliacao_completa.xlsx'})
            elif rev.mode == 'simples':
                downloads.append({'label': 'Doc Revisado (Simples)', 'file': f'{base_url}{nome_base}_revisao_simples.docx'})
                downloads.append({'label': 'Relatório Técnico', 'file': f'{base_url}relatorio_tecnico_simples_{nome_base}.docx'})
                downloads.append({'label': 'Planilha de Análise', 'file': f'{base_url}relatorio_revisao_simples.xlsx'})
            elif rev.mode == 'inconsistencias':
                # Não há 'Documento Revisado' neste modo, apenas relatórios.
                downloads.append({'label': 'Relatório de Inconsistências', 'file': f'{base_url}relatorio_tecnico_inconsistencias_{nome_base}.docx'})
                downloads.append({'label': 'Planilha de Análise', 'file': f'{base_url}relatorio_inconsistencias_{nome_base}.xlsx'})
        
        historico_list.append({
            'id': rev.id,
            'timestamp': rev.created_at,
            'file_name': rev.file_name,
            'mode': rev.mode.capitalize(),
            'status': rev.status,
            'status_display': rev.get_status_display(), # Para exibir o status formatado
            'downloads': downloads
        })

    context = {
        'revisoes': historico_list
    }
    return render(request, 'revisor/historico.html', context)